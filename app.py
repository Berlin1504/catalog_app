"""Product Catalog Flask App.

A small web app to build a veterinary product catalog in DOCX format,
matching the original template's banner + 5-column table layout.

Run:
    pip install -r requirements.txt
    python app.py
Then open http://localhost:5000 in your browser.
"""

from __future__ import annotations

import json
import uuid
import os
from pathlib import Path

from flask import (
    Flask,
    flash,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    url_for,
)
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor


# ----------------------------------------------------------------------------
# Paths & Flask config
# ----------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent

UPLOAD_FOLDER = BASE_DIR / "uploads"
OUTPUT_FOLDER = BASE_DIR / "output"
LEGACY_DATA_FILE = BASE_DIR / "data.json"

for d in (UPLOAD_FOLDER, OUTPUT_FOLDER):
    d.mkdir(parents=True, exist_ok=True)

app = Flask(__name__)
app.secret_key = "catalog-app-secret-change-me-in-production"
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["OUTPUT_FOLDER"] = str(OUTPUT_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB upload cap

# Database configuration
app.config["SQLALCHEMY_DATABASE_URI"] = f"sqlite:///{BASE_DIR / 'catalog.db'}"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

ALLOWED_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp"}

# ----------------------------------------------------------------------------
# Database Models
# ----------------------------------------------------------------------------
class Banner(db.Model):
    id = db.Column(db.String(8), primary_key=True)
    name = db.Column(db.String(255), nullable=False)

class Product(db.Model):
    id = db.Column(db.String(8), primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    composition = db.Column(db.Text, default="[]")  # JSON encoded list
    indications = db.Column(db.Text, default="[]")  # JSON encoded list
    features = db.Column(db.Text, default="[]")     # JSON encoded list
    image = db.Column(db.String(255), nullable=True)
    banner_id = db.Column(db.String(8), db.ForeignKey('banner.id'), nullable=True)
    
    banner = db.relationship('Banner', backref=db.backref('products', lazy=True, cascade="all, delete-orphan"))

    @property
    def composition_list(self):
        return json.loads(self.composition) if self.composition else []

    @property
    def indications_list(self):
        return json.loads(self.indications) if self.indications else []

    @property
    def features_list(self):
        return json.loads(self.features) if self.features else []

# ----------------------------------------------------------------------------
# Initialization & Migration
# ----------------------------------------------------------------------------
def init_db():
    with app.app_context():
        db.create_all()
        
        # Migrate from legacy data.json if it exists and DB is empty
        if LEGACY_DATA_FILE.exists() and Banner.query.count() == 0:
            print("Migrating data from legacy data.json to SQLite database...")
            try:
                data = json.loads(LEGACY_DATA_FILE.read_text(encoding="utf-8"))
                
                # Migrate banners
                banners_data = data.get("banners", [])
                if not banners_data and ("banner_1" in data or "banner_2" in data):
                    banners_data = [
                        {"id": "1", "name": data.get("banner_1", "خط الإنتاج الأول")},
                        {"id": "2", "name": data.get("banner_2", "خط الإنتاج الثاني")}
                    ]
                
                for b_data in banners_data:
                    if not Banner.query.get(b_data["id"]):
                        banner = Banner(id=b_data["id"], name=b_data["name"])
                        db.session.add(banner)
                
                # Migrate products
                for p_data in data.get("products", []):
                    if not Product.query.get(p_data["id"]):
                        product = Product(
                            id=p_data["id"],
                            name=p_data["name"],
                            composition=json.dumps(p_data.get("composition", []), ensure_ascii=False),
                            indications=json.dumps(p_data.get("indications", []), ensure_ascii=False),
                            features=json.dumps(p_data.get("features", []), ensure_ascii=False),
                            image=p_data.get("image"),
                            banner_id=p_data.get("banner", "1")
                        )
                        db.session.add(product)
                        
                db.session.commit()
                print("Migration complete. You can safely ignore data.json now.")
            except Exception as e:
                print(f"Migration error: {e}")
                db.session.rollback()

        # If no banners exist at all, create defaults
        if Banner.query.count() == 0:
            db.session.add(Banner(id="1", name="خط الإنتاج الأول"))
            db.session.add(Banner(id="2", name="خط الإنتاج الثاني"))
            db.session.commit()

init_db()

# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTS

def split_lines(text: str) -> list[str]:
    return [line.strip() for line in (text or "").splitlines() if line.strip()]

# ----------------------------------------------------------------------------
# Routes
# ----------------------------------------------------------------------------
@app.route("/")
def index():
    banners = Banner.query.all()
    products = Product.query.all()
    # Mock data structure to match what index.html expects for now
    data = {
        "banners": [{"id": b.id, "name": b.name} for b in banners],
        "products": products
    }
    return render_template("index.html", data=data)


@app.route("/update_banners", methods=["POST"])
def update_banners():
    banners = Banner.query.all()
    for banner in banners:
        new_name = request.form.get(f"banner_name_{banner.id}")
        if new_name and new_name.strip():
            banner.name = new_name.strip()
            
    new_banner_name = request.form.get("new_banner_name", "").strip()
    if new_banner_name:
        new_banner = Banner(id=uuid.uuid4().hex[:8], name=new_banner_name)
        db.session.add(new_banner)
        
    db.session.commit()
    flash("تم حفظ أسماء الخطوط", "success")
    return redirect(url_for("index"))


@app.route("/delete_banner/<banner_id>")
def delete_banner(banner_id: str):
    banner = Banner.query.get(banner_id)
    if banner:
        db.session.delete(banner)
        db.session.commit()
        flash("تم حذف الخط", "success")
    return redirect(url_for("index"))


@app.route("/add", methods=["POST"])
def add_product():
    name = request.form.get("name", "").strip()
    if not name:
        flash("اسم المنتج مطلوب", "error")
        return redirect(url_for("index"))

    composition = split_lines(request.form.get("composition", ""))
    indications = split_lines(request.form.get("indications", ""))
    features = split_lines(request.form.get("features", ""))
    banner_id = request.form.get("banner", "1")

    image_filename = None
    if "image" in request.files:
        f = request.files["image"]
        if f and f.filename and allowed_file(f.filename):
            ext = f.filename.rsplit(".", 1)[1].lower()
            image_filename = f"{uuid.uuid4().hex}.{ext}"
            safe_name = secure_filename(image_filename)
            f.save(UPLOAD_FOLDER / safe_name)
            image_filename = safe_name

    product = Product(
        id=uuid.uuid4().hex[:8],
        name=name,
        composition=json.dumps(composition, ensure_ascii=False),
        indications=json.dumps(indications, ensure_ascii=False),
        features=json.dumps(features, ensure_ascii=False),
        banner_id=banner_id,
        image=image_filename
    )
    db.session.add(product)
    db.session.commit()
    
    flash(f"تمت إضافة المنتج: {name}", "success")
    return redirect(url_for("index"))


@app.route("/edit/<product_id>", methods=["GET", "POST"])
def edit_product(product_id: str):
    product = Product.query.get_or_404(product_id)
    
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        if not name:
            flash("اسم المنتج مطلوب", "error")
            return redirect(url_for("edit_product", product_id=product_id))

        product.name = name
        product.composition = json.dumps(split_lines(request.form.get("composition", "")), ensure_ascii=False)
        product.indications = json.dumps(split_lines(request.form.get("indications", "")), ensure_ascii=False)
        product.features = json.dumps(split_lines(request.form.get("features", "")), ensure_ascii=False)
        product.banner_id = request.form.get("banner")

        if "image" in request.files:
            f = request.files["image"]
            if f and f.filename and allowed_file(f.filename):
                if product.image:
                    try:
                        (UPLOAD_FOLDER / product.image).unlink(missing_ok=True)
                    except Exception:
                        pass
                
                ext = f.filename.rsplit(".", 1)[1].lower()
                image_filename = f"{uuid.uuid4().hex}.{ext}"
                safe_name = secure_filename(image_filename)
                f.save(UPLOAD_FOLDER / safe_name)
                product.image = safe_name

        db.session.commit()
        flash(f"تم تحديث المنتج: {name}", "success")
        return redirect(url_for("index"))
        
    banners = Banner.query.all()
    return render_template("edit.html", product=product, banners=banners)


@app.route("/delete/<product_id>")
def delete_product(product_id: str):
    product = Product.query.get(product_id)
    if product:
        if product.image:
            try:
                (UPLOAD_FOLDER / product.image).unlink(missing_ok=True)
            except Exception:
                pass
        db.session.delete(product)
        db.session.commit()
        flash("تم حذف المنتج", "success")
    return redirect(url_for("index"))


@app.route("/clear")
def clear_all():
    Product.query.delete()
    Banner.query.delete()
    db.session.commit()
    
    db.session.add(Banner(id="1", name="خط الإنتاج الأول"))
    db.session.add(Banner(id="2", name="خط الإنتاج الثاني"))
    db.session.commit()

    for f in UPLOAD_FOLDER.iterdir():
        try:
            f.unlink()
        except Exception:
            pass
            
    flash("تم مسح كل البيانات", "success")
    return redirect(url_for("index"))


@app.route("/uploads/<path:filename>")
def uploaded_file(filename: str):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


@app.route("/generate")
def generate():
    products = Product.query.all()
    if not products:
        flash("لا توجد منتجات لتوليد الكتالوج", "error")
        return redirect(url_for("index"))

    out_name = "product_catalog.docx"
    out_path = OUTPUT_FOLDER / out_name

    build_catalog(products, out_path)

    return send_file(
        out_path,
        as_attachment=True,
        download_name=out_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ----------------------------------------------------------------------------
# DOCX generation
# ----------------------------------------------------------------------------
COL_WIDTHS_DXA = [3685, 3402, 2948, 2268, 3402]
COL_HEADERS = ["المميزات", "دواعي الاستخدام", "التركيب", "العبوة", "اسم المنتج"]
HEADER_FILL = "FFFF00"      # yellow
BANNER_FILL = "8BC53F"      # green
NAME_FILL = "FFFF00"        # yellow
RED_NUM = "C00000"          # red number prefix


def dxa_to_emu(dxa: int) -> int:
    return int(dxa * 12700 / 20)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_fill(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=80, left=80, bottom=80, right=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)


def _set_rtl_paragraph(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def _set_run_rtl(run) -> None:
    r_pr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    r_pr.append(rtl)


def _set_run_color(run, color_hex: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    r_pr.append(color)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tbl_pr.append(borders)


def _set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def _set_section_landscape(section) -> None:
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)


def _set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Emu(dxa_to_emu(width_dxa))


def _clear_cell_paragraphs(cell) -> None:
    for p in list(cell.paragraphs):
        cell._tc.remove(p._p)


def _add_rtl_text_paragraph(cell, text: str, *, bold=False, size_pt=None,
                            color_hex="000000", align=WD_ALIGN_PARAGRAPH.RIGHT,
                            spacing_after=0):
    p = cell.add_paragraph()
    p.alignment = align
    _set_rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    _set_run_color(run, color_hex)
    _set_run_rtl(run)
    return p


def _add_banner(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_layout_fixed(table)
    _set_table_borders(table)

    grid = table._tbl.find(qn("w:tblGrid"))
    grid.clear()
    gc = OxmlElement("w:gridCol")
    gc.set(qn("w:w"), "15705")
    grid.append(gc)

    cell = table.cell(0, 0)
    _set_cell_borders(cell)
    _set_cell_fill(cell, BANNER_FILL)
    _set_cell_margins(cell, top=140, left=120, bottom=140, right=120)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_width(cell, 15704)
    _clear_cell_paragraphs(cell)

    _add_rtl_text_paragraph(cell, text or "—",
                             bold=True, size_pt=20,
                             color_hex="000000",
                             align=WD_ALIGN_PARAGRAPH.CENTER)

    spacer = doc.add_paragraph()
    _set_rtl_paragraph(spacer)


def _new_product_table(doc):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    _set_table_layout_fixed(table)
    _set_table_borders(table)

    grid = table._tbl.find(qn("w:tblGrid"))
    grid.clear()
    for w in COL_WIDTHS_DXA:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)

    return table


def _fill_header_row(table) -> None:
    row = table.rows[0]
    for i, header in enumerate(COL_HEADERS):
        cell = row.cells[i]
        _set_cell_borders(cell)
        _set_cell_fill(cell, HEADER_FILL)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_width(cell, COL_WIDTHS_DXA[i])
        _clear_cell_paragraphs(cell)
        _add_rtl_text_paragraph(cell, header,
                                 bold=True, size_pt=14,
                                 color_hex="000000",
                                 align=WD_ALIGN_PARAGRAPH.CENTER)


def _fill_product_row(table, product, image_path: str | None) -> None:
    row = table.add_row()
    cells = row.cells
    for i in range(5):
        _set_cell_width(cells[i], COL_WIDTHS_DXA[i])

    # 1. Features
    c = cells[0]
    _set_cell_borders(c)
    _set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_paragraphs(c)
    features = product.features_list or ["—"]
    for idx, feat in enumerate(features, start=1):
        p = c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        _set_rtl_paragraph(p)
        r1 = p.add_run(f"{idx}. ")
        r1.bold = True
        _set_run_color(r1, RED_NUM)
        _set_run_rtl(r1)
        r2 = p.add_run(feat)
        _set_run_color(r2, "000000")
        _set_run_rtl(r2)

    # 2. Indications
    c = cells[1]
    _set_cell_borders(c)
    _set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_paragraphs(c)
    indications = product.indications_list or ["—"]
    for ind in indications:
        _add_rtl_text_paragraph(c, ind, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # 3. Composition
    c = cells[2]
    _set_cell_borders(c)
    _set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_paragraphs(c)
    comp = product.composition_list or ["—"]
    for line in comp:
        _add_rtl_text_paragraph(c, line, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # 4. Image
    c = cells[3]
    _set_cell_borders(c)
    _set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_paragraphs(c)
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_paragraph(p)
    if image_path and Path(image_path).exists():
        try:
            run = p.add_run()
            run.add_picture(image_path, width=Inches(1.5))
        except Exception:
            _add_rtl_text_paragraph(c, "[صورة المنتج]",
                                    align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _add_rtl_text_paragraph(c, "[صورة المنتج]",
                                align=WD_ALIGN_PARAGRAPH.CENTER)

    # 5. Name
    c = cells[4]
    _set_cell_borders(c)
    _set_cell_fill(c, NAME_FILL)
    _set_cell_margins(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_paragraphs(c)
    _add_rtl_text_paragraph(c, product.name or "—",
                            bold=True, size_pt=14,
                            color_hex="000000",
                            align=WD_ALIGN_PARAGRAPH.CENTER)


def build_catalog(products: list, output_path: Path) -> None:
    doc = Document()
    _set_section_landscape(doc.sections[0])

    banners = Banner.query.all()
    products_by_banner = {}
    for p in products:
        products_by_banner.setdefault(p.banner_id, []).append(p)

    first = True
    for banner in banners:
        banner_products = products_by_banner.get(banner.id, [])
        if not banner_products:
            continue

        if not first:
            _add_page_break(doc)
        first = False

        _add_banner(doc, banner.name)
        table = _new_product_table(doc)
        _fill_header_row(table)
        for product in banner_products:
            image_path = None
            if product.image:
                cand = UPLOAD_FOLDER / product.image
                if cand.exists():
                    image_path = str(cand)
            _fill_product_row(table, product, image_path)

    doc.save(output_path)


# ----------------------------------------------------------------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)